from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import pandas as pd
from docx.enum.table import WD_ALIGN_VERTICAL
from pathlib import Path


class CourseReport:

	def __init__(self, template,folder,summary_data,attainment_data,course,level,gender,simd,asn):
		self.gender = gender
		self.template = template
		self.summary_data = summary_data
		self.attainment_data = attainment_data
		self.doc = Document(self.template)
		self.course = course
		self.simd = simd
		self.asn = asn
		self.folder = folder
	
		self.levels = {}
		self.levels[74] = "National 4"
		self.levels[75] = "National 5"
		self.levels[76] = "Higher"
		self.levels[77] = "Advanced Higher"
				
		self.level = self.levels[level]
		self.fname = self.course+"-"+self.level+".docx"
		
		self.create_doc(self.course,self.level)
		self.save()
		
	def create_doc(self,course, level):
		section = self.doc.sections[0]
		header = section.header

		table = header.tables[0]
		table.cell(0,1).text = table.cell(0,1).text + course
		table.cell(0,2).text = table.cell(0,2).text + level
		
		table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		table.cell(0,2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		
		
		self.doc.tables #a list of all tables in document
		table = self.doc.tables[0]
		stages = ["S4","S5","S6"]
		for i in range(2,5):
			temp = self.summary_data[stages[i-2]]
			
			if (isinstance(temp, dict)):
				temp_ccv = self.attainment_data[stages[i-2]]
				table.cell(i,1).add_table(rows=2, cols=6)
				table.cell(i,1).add_table(rows=3, cols=6)		
				
				new_table = table.cell(i,1).tables[0]
				new_table.style = 'Table Grid'
				self.add_summary(new_table,temp,temp_ccv)
				self.change_font_size(new_table,8)
			
				new_table = table.cell(i,1).tables[1]
				new_table.style = 'Table Grid'
				self.add_results(new_table,temp,temp_ccv)
				self.change_font_size(new_table,8)
			else:
				table.cell(i,1).add_table(rows=1, cols=1)
				table.cell(i,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				new_table = table.cell(i,1).tables[0]
				new_table.cell(0,0).text = "No data"
			
			
# 		SIMD	
		table = self.doc.tables[1]
		table.cell(1,1).add_table(rows=6, cols=7)
		new_table = table.cell(1,1).tables[0]
		new_table.style = 'Table Grid'
				
		self.add_simd(new_table,"data")
		self.change_font_size(new_table,8)
		
		
		
# 		Gender
		table = self.doc.tables[1]
		table.cell(3,1).add_table(rows=5, cols=7)
		new_table = table.cell(3,1).tables[0]
		new_table.style = 'Table Grid'
				
		self.add_gender(new_table,"data")
		self.change_font_size(new_table,8)
		
# ASN
		table = self.doc.tables[1]
		table.cell(5,1).add_table(rows=4, cols=7)
		new_table = table.cell(5,1).tables[0]
		new_table.style = 'Table Grid'
		self.add_asn(new_table,"data")
		self.change_font_size(new_table,8)
		

	def add_results(self,table,data,ccv):
		
		school = data["school"]
		national = data["national"]
		
		table.cell(1,0).text = "Portree"
		table.cell(2,0).text = "National"

		table.cell(0,1).text = "A"
		table.cell(0,2).text = "A-B %"
		table.cell(0,3).text = "A-C %"
		table.cell(0,4).text = "A-D %"
		table.cell(0,5).text = "NA %"
		
		school_disp = [""]*5
		national_disp = [""]*5
		

		school_disp[0] = str(school['% Grade A'].values[0])
		school_disp[1] = str(school['% Grades A to B'].values[0])
		school_disp[2] = str(school['% Grades A to C'].values[0])
		school_disp[3] = str(school['% Grades A to D'].values[0])
		school_disp[4] = str(school['% No Award'].values[0])
		
		national_disp[0] = str(national['National % Grade A'].values[0])
		national_disp[1] = str(national['National % Grades A to B'].values[0])
		national_disp[2] = str(national['National % Grades A to C'].values[0])
		national_disp[3] = str(national['National % Grades A to D'].values[0])
		national_disp[4] = str(national['National % No Award'].values[0])
		
		for i in range(1,6):
		# 	School
			table.cell(1,i).text = str(school_disp[i-1])
		# National
			table.cell(2,i).text = str(national_disp[i-1])
			table.cell(1,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			table.cell(2,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
	
		
		
	def change_font_size(self,table,size):		
		for row in table.rows:
			for cell in row.cells:
				paragraphs = cell.paragraphs
				for paragraph in paragraphs:
					for run in paragraph.runs:
						font = run.font
						font.size= Pt(size)
		
	def add_simd(self,table,data):
		
		table.cell(1,0).text = "1"
		table.cell(2,0).text = "2"
		table.cell(3,0).text = "3"
		table.cell(4,0).text = "4"
		table.cell(5,0).text = "5"

		table.cell(0,1).text = "A%"
		table.cell(0,2).text = "B%"
		table.cell(0,3).text = "C%"
		table.cell(0,4).text = "D%"
		table.cell(0,5).text = "NA%"
		table.cell(0,6).text = "Num Entries"
		
		for i in range(1,7):
			table.cell(0,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		for i in range(1,6):
			table.cell(i,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
		
		
		for i in range(1,6):
			temp = self.simd[i]
# 			
			school = temp[temp["Establishment"]=="Portree High School"]
			national = temp[temp["Establishment"]=="The National Establishment"] 
			if (school.size!=0):
				count = 1
				for grade in ["% Grade A","% Grade B","% Grade C","% Grade D","% No Award"]:
					table.cell(i,count).text = str(school[grade].values[0])
					table.cell(i,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
					count +=1
				entries = school["# Grade A"]+school["# Grade B"]+school["# Grade C"]+school["# Grade D"]+school["# No Award"]
				table.cell(i,count).text = str(int(entries.values[0]))
				table.cell(i,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
				
	
	
	def add_gender(self,table,data):
# 		school = data["school"]
		female = self.gender["female"]
		male = self.gender["male"]
		
		table.cell(1,0).text = "Female (School)"
		table.cell(2,0).text = "Male (School)"
		table.cell(3,0).text = "Female (National)"
		table.cell(4,0).text = "Male (National)"
		for i in range(1,5):
			table.cell(i,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
		count = 1
		for i in ["A%","B%","C%","D%","NA%","Num Entries"]:
			table.cell(0,count).text = i
			table.cell(0,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
			count +=1
		
		row_count = 1
		for group in ["Portree High School","The National Establishment"]:
			total_entries = 0.0  

			temp = female[female["Establishment"]==group]
			entries = temp[["# Grade A","# Grade B","# Grade C","# Grade D","# No Award"]]
			total_entries = entries.values.sum()
			
			temp = male[male["Establishment"]==group]
			entries = temp[["# Grade A","# Grade B","# Grade C","# Grade D","# No Award"]]
			total_entries += entries.values.sum()



			school = female[female["Establishment"]==group]
			if (school.size!=0):
				count = 1
				for grade in ["% Grade A","% Grade B","% Grade C","% Grade D","% No Award"]:
					table.cell(row_count,count).text = str(school[grade].values[0])
					table.cell(row_count,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
					count +=1
				entries = school[["# Grade A","# Grade B","# Grade C","# Grade D","# No Award"]]
				entries = entries.values.sum()
				table.cell(row_count,count).text = str(int(entries))+" ("+ str(round(entries/total_entries*100,0)) +"%)"
				table.cell(row_count,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
				
			row_count +=1
			school = male[male["Establishment"]==group]
			if (school.size!=0):
				count = 1
				for grade in ["% Grade A","% Grade B","% Grade C","% Grade D","% No Award"]:
					table.cell(row_count,count).text = str(school[grade].values[0])
					table.cell(row_count,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		
					count +=1
				entries = school[["# Grade A","# Grade B","# Grade C","# Grade D","# No Award"]]
				entries = entries.values.sum()
				table.cell(row_count,count).text = str(int(entries))+" ("+ str(round(entries/total_entries*100,0)) +"%)"
				table.cell(row_count,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
				
			row_count +=1	
		
	def add_asn(self,table,data):
# 		school = data["school"]
		
		table.cell(1,0).text = "Portree"
		table.cell(2,0).text = "Highland"
		table.cell(3,0).text = "National"
		
		count = 1
		for i in ["A%","B%","C%","D%","NA%","Num Entries"]:
			table.cell(0,count).text = i
			table.cell(0,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
			count +=1
		
		row_count = 1
		for group in ["Portree High School","Highland","The National Establishment"]:
			
			school = self.asn[self.asn["Establishment"]==group]
			if (school.size!=0):
				count = 1
				for grade in ["% Grade A","% Grade B","% Grade C","% Grade D","% No Award"]:
					table.cell(row_count,count).text = str(school[grade].values[0])
					table.cell(row_count,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		
					count +=1
				entries = school[["# Grade A","# Grade B","# Grade C","# Grade D","# No Award"]]
				entries = entries.values.sum()
				table.cell(row_count,count).text = str(int(entries))
				table.cell(row_count,count).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			
			row_count +=1
	
	
	def add_summary(self,table,data,ccv):
		table.cell(1,0).text = "Portree"

		table.cell(0,1).text = "Entries"
		table.cell(0,2).text = "Pass Rate"
		table.cell(0,3).text = "% Base Cohort"
		table.cell(0,4).text = "CCV"
		table.cell(0,5).text = "CCV Sig"
		
		for i in range(1,6):
			table.cell(0,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		
		
		
		school = data["school"]
		
		if(isinstance(ccv,pd.DataFrame)):
			if(ccv.size>0):
				temp = ccv[["ccv","sig"]].values[0]
				ccv = temp[0]
				sig = temp[1]
			else:
				ccv = "No data"
				sig = "No data"
		else:
				ccv = "No data"
				sig = "No data"
		
			
		school_disp = [""]*5


		school_disp[0] = str(school['Total'].values[0])
		school_disp[1] = str(school['% Grades A to D'].values[0])
		school_disp[2] = str(school['Resulted Entries % of Base Cohort'].values[0])
		school_disp[3] = str(ccv)
		school_disp[4] = str(sig)
	
		for i in range(1,6):
		# 	School
			table.cell(1,i).text = str(school_disp[i-1])
			table.cell(1,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		
			
	def save(self):
		self.doc.save(self.folder/self.fname)
		


		
	